import requests
import yfinance as yf
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import time

def fetch_company_info(ipo_name):
    try:
        search_url = f"https://query1.finance.yahoo.com/v1/finance/search?q={ipo_name}"
        response = requests.get(search_url)

        if response.status_code == 429:
            print("Too many requests. Please wait before trying again.")
            time.sleep(120)  # Wait for 60 seconds before retrying
            return fetch_company_info(ipo_name)  # Retry the request

        if response.status_code != 200:
            print(f"Error: Received response with status code {response.status_code}")
            print(f"Response content: {response.text}")
            return None, None

        results = response.json().get('quotes', [])

        if results:
            company_data = results[0]  # Get the first match
            ticker = company_data['symbol']
            sector = company_data['sector']
            return ticker, sector
        else:
            print("No results found for the given IPO name.")
            return None, None
    except requests.exceptions.RequestException as e:
        print(f"Request failed: {e}")
        return None, None
    except ValueError as ve:
        print(f"JSON decoding failed: {ve}")
        return None, None


def fetch_nearest_competitor(ticker, sector):
    if ticker is None:
        print("No ticker provided; cannot fetch nearest competitor.")
        return None

    try:
        company = yf.Ticker(ticker)
        competitors = company.recommendations  # You can adjust this based on your needs
        if competitors is not None and not competitors.empty:
            competitor_tickers = competitors['ticker'].unique()  # Extract unique competitor tickers
            return competitor_tickers[:5]  # Return the first 5 competitors for simplicity
        else:
            print(f"No competitors found for {ticker}.")
            return None
    except Exception as e:
        print(f"Failed to fetch competitors for {ticker}: {e}")
        return None


def analyze_profitability(document_path, ipo_name):
    ticker, sector = fetch_company_info(ipo_name)
    if ticker:
        competitors = fetch_nearest_competitor(ticker, sector)
        if competitors is None:
            return  # Exit if no competitors found

        # Fetch financial data for the IPO company
        company = yf.Ticker(ticker)
        stock_info = company.info

        if stock_info:
            market_cap = stock_info.get('marketCap', 'N/A')
            earnings = stock_info.get('earningsQuarterly', 'N/A')
            # Add more financial metrics as needed

        # Prepare the document
        document = Document()
        document.add_heading(f'{ipo_name} Financial Analysis', level=1)
        document.add_paragraph(f'Ticker: {ticker}')
        document.add_paragraph(f'Sector: {sector}')
        document.add_paragraph(f'Market Capitalization: {market_cap}')

        # Pie chart for market share (placeholder data for demonstration)
        sizes = [70, 20, 10]  # Placeholder sizes for pie chart
        labels = [ticker] + list(competitors)
        plt.figure(figsize=(8, 6))
        plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
        plt.axis('equal')  # Equal aspect ratio ensures that pie chart is a circle.
        plt.title(f'Market Share Analysis for {ipo_name}')
        plt.savefig('market_share_analysis.png')  # Save the figure
        plt.close()

        # Add pie chart to document
        document.add_picture('market_share_analysis.png')

        # Save the document
        report_filename = f'{ipo_name}_IPO_Analysis_Report.docx'
        document.save(report_filename)
        print(f'Report saved as {report_filename}')
    else:
        print("No valid ticker found. Exiting analysis.")


# Example usage
document_path = "C:\\Users\\shrav\\Downloads\\Swiggy Limited Prospectus.pdf"  # Replace with your PDF document path
ipo_name = "Swiggy Limited"  # Replace with the IPO name you are analyzing
analyze_profitability(document_path, ipo_name)
