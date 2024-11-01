import requests
import fitz  # PyMuPDF for PDF extraction
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from textblob import TextBlob
import re
from docx import Document
import io


# Function to extract text from PDF document
def extract_text_from_pdf(pdf_path):
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text()
    return text


# Function to search for financial keywords in document
def search_keywords(text, keywords):
    found = {}
    for keyword in keywords:
        matches = re.findall(rf"\b{keyword}\b", text, re.IGNORECASE)
        found[keyword] = len(matches)
    return found


# Function to get news sentiment using NewsAPI
def get_news_sentiment(query):
    api_key = "6f83f736660a419bb76d73249e58b264"  # Replace with your NewsAPI key
    url = f"https://newsapi.org/v2/everything?q={query}&apiKey={api_key}"
    response = requests.get(url)
    articles = response.json().get("articles", [])

    headlines = [article["title"] for article in articles]
    sentiments = [TextBlob(headline).sentiment.polarity for headline in headlines]

    avg_sentiment = np.mean(sentiments) if sentiments else 0
    return avg_sentiment


# Analyze profitability in-depth based on document data and sentiment
def analyze_profitability(document_path, ipo_name):
    # Step 1: Extract and analyze document
    keywords = ['revenue', 'growth', 'debt', 'profit', 'competitive', 'cash flow', 'risk']
    document_text = extract_text_from_pdf(document_path)
    keyword_counts = search_keywords(document_text, keywords)

    # Step 2: Get sentiment score for recent news
    news_sentiment = get_news_sentiment(ipo_name)

    # Step 3: Compile data into a DataFrame
    data = pd.DataFrame(list(keyword_counts.items()), columns=['Keyword', 'Count'])
    data['Sentiment'] = news_sentiment
    data['Impact'] = data['Count'] * (1 + data['Sentiment'])

    # Step 4: Visualization
    plt.figure(figsize=(10, 6))
    sns.barplot(x='Keyword', y='Impact', data=data, palette='viridis')
    plt.title(f"Impact Score Analysis for IPO: {ipo_name}")
    plt.xlabel('Financial Factor')
    plt.ylabel('Impact Score')
    graph_image = io.BytesIO()
    plt.savefig(graph_image, format='png')
    plt.close()

    # Step 5: Generate report
    document = Document()
    document.add_heading(f'Detailed Analysis Report for IPO: {ipo_name}', level=1)

    document.add_paragraph("**1. Financial Factor Analysis:**")
    for idx, row in data.iterrows():
        keyword = row['Keyword']
        impact = row['Impact']
        if impact > 10:
            reason = f"{keyword.capitalize()} is strong, indicating financial stability and growth potential."
        else:
            reason = f"Low {keyword} may pose risks to profitability due to weaker financial indicators."
        document.add_paragraph(f"{keyword.capitalize()}: Impact Score = {impact:.2f}. Reason: {reason}")

    document.add_paragraph("**2. Sentiment and Market Position Analysis:**")
    if news_sentiment > 0:
        sentiment_analysis = f"Positive market sentiment ({news_sentiment:.2f}), indicating potential market support."
    else:
        sentiment_analysis = f"Negative market sentiment ({news_sentiment:.2f}), indicating skepticism or competition concerns."
    document.add_paragraph(sentiment_analysis)

    # Step 6: Conclusion
    avg_impact = data['Impact'].mean()
    conclusion = "IPO has a high potential for profitability." if avg_impact > 10 else "IPO might be risky or less profitable."
    document.add_paragraph("**Conclusion:**")
    document.add_paragraph(conclusion)

    # Add Graph to the report
    document.add_paragraph("**Impact Score Analysis Graph:**")
    graph_image.seek(0)
    document.add_picture(graph_image, width=5000000)  # adjust width if necessary

    # Save the report
    report_filename = f"{ipo_name}_IPO_Analysis_Report.docx"
    document.save(report_filename)
    print(f"Report generated: {report_filename}")

document_path = 'C:\\Users\\shrav\\Downloads\\Swiggy Limited Prospectus.pdf'
ipo_name = 'Swiggy Limited'
analyze_profitability(document_path, ipo_name)